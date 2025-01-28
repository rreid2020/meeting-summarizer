# export_service.py
from fastapi.responses import FileResponse
from docx import Document
import json
from datetime import datetime

class ExportService:
    async def export_summary(self, summary: dict, format: str):
        exporters = {
            'txt': self._export_txt,
            'pdf': self._export_pdf,
            'docx': self._export_docx,
            'json': self._export_json
        }
        
        if format not in exporters:
            raise HTTPException(400, f"Unsupported format: {format}")
            
        return await exporters[format](summary)

    async def _export_docx(self, summary):
        doc = Document()
        doc.add_heading('Meeting Summary', 0)
        doc.add_paragraph(summary['summary'])
        
        doc.add_heading('Action Items', level=1)
        for item in summary['action_items']:
            doc.add_paragraph(f"• {item['action']}", style='List Bullet')
        
        filename = f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        return FileResponse(filename)